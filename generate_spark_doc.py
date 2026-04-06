from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.style = "No Spacing"
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def generate_doc():
    doc = Document()

    # Title
    title = doc.add_heading("Top 100 Spark Differences - Part 1 (1-10)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    differences = [
        {
            "id": 1,
            "title": "RDD vs. DataFrame",
            "theory": "RDD is Spark's fundamental data abstraction, providing low-level, schema-less distributed collections. DataFrames are high-level structured APIs with a schema (columns/types), optimized by the Catalyst optimizer for better performance.",
            "pyspark": "rdd = sc.parallelize([(1, 'A'), (2, 'B')])\ndf = spark.createDataFrame(rdd, ['id', 'name'])",
            "pandas": "import pandas as pd\ndf = pd.DataFrame({'id': [1, 2], 'name': ['A', 'B']})",
            "tip": "Use DataFrames/Datasets for most tasks; reserve RDDs for low-level processing where Catalyst cannot optimize.",
        },
        {
            "id": 2,
            "title": "Transformation vs. Action",
            "theory": "Transformations are lazy; they build a Logical Plan but don't execute until an action is called. Actions are eager; they trigger the computation and return a result to the driver or write data to storage.",
            "pyspark": "# Transformation (Lazy)\ndf_filtered = df.filter(df['id'] > 1)\n\n# Action (Eager)\ndf_filtered.show()",
            "pandas": "# Pandas is eager by default\ndf_filtered = df[df['id'] > 1]  # Immediate result",
            "tip": "Chain transformations to minimize shuffles; use actions sparingly to avoid unnecessary cluster activity.",
        },
        {
            "id": 3,
            "title": "Narrow vs. Wide Transformation",
            "theory": "Narrow transformations (e.g., map, filter) do not require data movement across nodes (1:1 partition mapping). Wide transformations (e.g., groupBy, join) require a 'shuffle' where data is redistributed across executors.",
            "pyspark": "# Narrow\ndf_map = df.select('id')\n\n# Wide (Triggers Shuffle)\ndf_grouped = df.groupBy('id').count()",
            "pandas": "# Local operations; no 'shuffle' concept but grouping is computationally heavier\ndf_grouped = df.groupby('id').size()",
            "tip": "Maximize narrow transformations; wide transformations are the most expensive part of Spark jobs.",
        },
        {
            "id": 4,
            "title": "Repartition vs. Coalesce",
            "theory": "Repartition performs a full shuffle to increase or decrease the number of partitions. Coalesce decreases partitions by merging them, avoiding a shuffle when possible, but cannot increase partitions.",
            "pyspark": "# Full shuffle\ndf_r = df.repartition(10)\n\n# Minimize shuffle (decrease only)\ndf_c = df.coalesce(2)",
            "pandas": "# Not applicable locally, but np.array_split is similar\nimport numpy as np\nsplit_dfs = np.array_split(df, 2)",
            "tip": "Use coalesce() when reducing partitions to improve performance; use repartition() to balance data skew.",
        },
        {
            "id": 5,
            "title": "Cache vs. Persist",
            "theory": "Cache() is a shorthand for Persist() using the default storage level (MEMORY_AND_DISK). Persist() allows custom storage levels (e.g., MEMORY_ONLY, DISK_ONLY, OFF_HEAP).",
            "pyspark": "from pyspark import StorageLevel\ndf.cache()\ndf.persist(StorageLevel.DISK_ONLY)",
            "pandas": "# In memory, df.copy() or storing in a variable persists the state locally\ndf_persisted = df.copy()",
            "tip": "Unpersist data once done to free up cluster memory; don't over-cache.",
        },
        {
            "id": 6,
            "title": "Partitioning vs. Bucketing",
            "theory": "Partitioning creates physical directories based on column values (good for high cardinality filter pushdown). Bucketing hashes data into a fixed number of files (buckets) within each partition (good for large table joins).",
            "pyspark": "# Save by date\ndf.write.partitionBy('date').parquet('path/')\n\n# Save into 10 buckets\ndf.write.bucketBy(10, 'id').saveAsTable('t')",
            "pandas": "# Partitioned write (using PyArrow)\ndf.to_parquet('path/', partition_cols=['date'])",
            "tip": "Partition on columns used in 'where' clauses; bucket on columns used in 'joins'.",
        },
        {
            "id": 7,
            "title": "Broadcast Join vs. Shuffle Hash Join",
            "theory": "Broadcast join sends the entire small table to every executor, avoiding a shuffle of the large table. Shuffle Hash Join redistributes both tables based on the join key (expensive shuffle).",
            "pyspark": "from pyspark.sql.functions import broadcast\ndf_large.join(broadcast(df_small), 'id')",
            "pandas": "# Always local in memory; pd.merge is equivalent\npd.merge(df1, df2, on='id')",
            "tip": "Use broadcast join if one table fits in executor memory (default limit is 10MB, configurable).",
        },
        {
            "id": 8,
            "title": "Map vs. MapPartitions",
            "theory": "Map applies a function to every row in an RDD. MapPartitions applies a function once per partition, allowing you to reuse expensive initializations (like DB connections) across all rows in that partition.",
            "pyspark": "# Row level\nrdd.map(lambda x: x * 2)\n\n# Partition level\nrdd.mapPartitions(lambda part: [sum(part)])",
            "pandas": "# df.apply is row/col level\ndf['id'].apply(lambda x: x * 2)",
            "tip": "Use mapPartitions for database connections or object initialization to reduce overhead.",
        },
        {
            "id": 9,
            "title": "GroupByKey vs. ReduceByKey",
            "theory": "GroupByKey shuffles all data to the same key first (high network overhead). ReduceByKey merges data locally before shuffling (Map-side combine), significantly reducing network traffic.",
            "pyspark": "# Heavy shuffle\nrdd.groupByKey()\n\n# Efficient (Map-side combine)\nrdd.reduceByKey(lambda a, b: a + b)",
            "pandas": "# .agg() is generally optimized for aggregation\ndf.groupby('key').sum()",
            "tip": "Prefer reduceByKey or aggregateByKey over groupByKey whenever possible.",
        },
        {
            "id": 10,
            "title": "Client Mode vs. Cluster Mode",
            "theory": "In Client mode, the Spark Driver runs on the machine where the job is submitted (useful for interactive shells). In Cluster mode, the Driver runs inside one of the worker nodes (safer for production jobs).",
            "pyspark": "# Client Mode CLI\nspark-submit --deploy-mode client script.py\n\n# Cluster Mode CLI\nspark-submit --deploy-mode cluster script.py",
            "pandas": "Not applicable; Pandas runs entirely on the local machine.",
            "tip": "Use Cluster mode for scheduled production workflows to prevent job failure if the local client disconnects.",
        },
    ]

    for item in differences:
        doc.add_heading(f"{item['id']}. {item['title']}", level=1)

        doc.add_heading("Theoretical Difference:", level=2)
        doc.add_paragraph(item["theory"])

        doc.add_heading("PySpark Implementation:", level=2)
        add_code_block(doc, item["pyspark"])

        doc.add_heading("Pandas Equivalent / Comparison:", level=2)
        add_code_block(doc, item["pandas"])

        doc.add_heading("Expert Tip:", level=2)
        tip_p = doc.add_paragraph()
        run = tip_p.add_run(item["tip"])
        run.bold = True

        doc.add_page_break()

    doc.save("Spark_Differences_Part1.docx")
    print("Document 'Spark_Differences_Part1.docx' generated successfully.")


if __name__ == "__main__":
    generate_doc()

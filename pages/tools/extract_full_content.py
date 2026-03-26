from docx import Document
import json

doc = Document('Pandas-SQL-Numpy-PySpark_CheatSheet.docx')

print("=" * 100)
print("FULL CONTENT FROM DOCX FILE")
print("=" * 100)

# Extract all paragraphs
print("\n=== PARAGRAPHS ===\n")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

# Extract all tables
print("\n\n=== TABLES ===\n")
for table_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {table_idx + 1} ---\n")
    
    # Get all rows
    for row_idx, row in enumerate(table.rows):
        cols = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cols))
        print("-" * 150)

print("\n" + "=" * 100)
print("END OF CONTENT")
print("=" * 100)

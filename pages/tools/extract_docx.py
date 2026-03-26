from docx import Document
import json

doc = Document('Pandas-SQL-Numpy-PySpark_CheatSheet.docx')

print("=" * 80)
print("PARAGRAPHS:")
print("=" * 80)
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n" + "=" * 80)
print("TABLES:")
print("=" * 80)
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ---")
    table_data = []
    for row in table.rows:
        cols = [cell.text.strip() for cell in row.cells]
        table_data.append(cols)
        print(" | ".join(cols))

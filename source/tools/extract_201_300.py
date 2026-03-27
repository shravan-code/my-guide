# -*- coding: utf-8 -*-
import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document

doc = Document("source/tools/SQL_Questions_201_to_300.docx")

questions_para = {}
current_num = None

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    import re

    q_match = re.match(r"^Q(\d+)\.\s*(.+)$", text)
    if q_match:
        current_num = int(q_match.group(1))
        questions_para[current_num] = {
            "question": q_match.group(2).strip(),
            "explanation": "",
        }
    elif text.startswith("📝") or text.startswith("Explanation:"):
        explanation = text.replace("📝", "").replace("Explanation:", "").strip()
        if current_num and current_num in questions_para:
            questions_para[current_num]["explanation"] = explanation

questions_data = []

for num in sorted(questions_para.keys()):
    q_data = {
        "number": num,
        "question": questions_para[num]["question"],
        "explanation": questions_para[num]["explanation"],
        "query": "",
        "input": [],
        "output": [],
    }
    questions_data.append(q_data)

table_idx = 0
for q_idx, q_data in enumerate(questions_data):
    if table_idx < len(doc.tables):
        table = doc.tables[table_idx]
        first_cell = (
            table.rows[0].cells[0].text.strip()
            if table.rows and table.rows[0].cells
            else ""
        )
        if "PostgreSQL Query" in first_cell or "🔷" in first_cell:
            query_lines = []
            for row in table.rows[1:]:
                for cell in row.cells:
                    if cell.text.strip():
                        query_lines.append(cell.text.strip())
            q_data["query"] = " ".join(query_lines).replace("\n", " ").strip()
    table_idx += 1

    if table_idx < len(doc.tables):
        table = doc.tables[table_idx]
        first_cell = (
            table.rows[0].cells[0].text.strip()
            if table.rows and table.rows[0].cells
            else ""
        )
        if "Input" in first_cell or "📥" in first_cell:
            header_text = (
                first_cell.replace("📥", "")
                .replace("Input:", "")
                .replace("Input", "")
                .strip()
            )
            headers = (
                [cell.text.strip() for cell in table.rows[1].cells]
                if len(table.rows) > 1
                else []
            )
            table_data = {"table": header_text, "headers": headers, "rows": []}
            for row in table.rows[2:]:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data["rows"].append(row_data)
            q_data["input"] = table_data
    table_idx += 1

    if table_idx < len(doc.tables):
        table = doc.tables[table_idx]
        first_cell = (
            table.rows[0].cells[0].text.strip()
            if table.rows and table.rows[0].cells
            else ""
        )
        if "Output" in first_cell or "📤" in first_cell:
            headers = (
                [cell.text.strip() for cell in table.rows[1].cells]
                if len(table.rows) > 1
                else []
            )
            table_data = {"headers": headers, "rows": []}
            for row in table.rows[2:]:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data["rows"].append(row_data)
            q_data["output"] = table_data
    table_idx += 1

import json

with open("source/sql/sql-questions-201-300.json", "w", encoding="utf-8") as f:
    json.dump(questions_data, f, indent=2, ensure_ascii=False)

print(f"Extracted {len(questions_data)} questions")

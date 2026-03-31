from pathlib import Path
import sqlite3

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
DOC_PATH = BASE_DIR / "Pandas-SQL-Numpy-PySpark_CheatSheet.docx"
DB_PATH = BASE_DIR / "scenarios.db"


def load_categories(document):
    categories = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text and paragraph.style.name == "Heading 1":
            categories.append(text.replace("  ", " - "))
    return categories


def recreate_database(connection):
    cursor = connection.cursor()
    cursor.execute("DROP TABLE IF EXISTS scenarios")
    cursor.execute(
        """
        CREATE TABLE scenarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            technology TEXT NOT NULL,
            category TEXT,
            scenario_name TEXT,
            description TEXT,
            code_snippet TEXT,
            example_output TEXT,
            raw_content TEXT,
            source_page INTEGER
        )
        """
    )
    connection.commit()


def normalize_header(value):
    cleaned = value.encode("ascii", "ignore").decode()
    cleaned = " ".join(cleaned.split())
    if cleaned == "PostgreSQL":
        return "PostgreSQL"
    if cleaned == "Pandas":
        return "Pandas"
    if cleaned == "NumPy":
        return "NumPy"
    if cleaned == "PySpark":
        return "PySpark"
    return cleaned


def extract_rows(document, categories):
    records = []

    for table_index, table in enumerate(document.tables):
        category = categories[table_index] if table_index < len(categories) else None
        headers = [normalize_header(cell.text.strip()) for cell in table.rows[0].cells]
        technologies = headers[1:]

        for row_index, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            scenario_name = cells[0]

            for technology, code_snippet in zip(technologies, cells[1:]):
                records.append(
                    (
                        technology,
                        category,
                        scenario_name,
                        None,
                        code_snippet,
                        None,
                        f"Table {table_index + 1} Row {row_index} | {technology}",
                        None,
                    )
                )

    return records


def main():
    document = Document(str(DOC_PATH))
    categories = load_categories(document)
    records = extract_rows(document, categories)

    connection = sqlite3.connect(DB_PATH)
    recreate_database(connection)

    connection.executemany(
        """
        INSERT INTO scenarios (
            technology,
            category,
            scenario_name,
            description,
            code_snippet,
            example_output,
            raw_content,
            source_page
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        records,
    )
    connection.commit()
    connection.close()

    print(f"Scenarios extracted and saved to {DB_PATH}")
    print(f"Total inserted rows: {len(records)}")


if __name__ == "__main__":
    main()
